"""
contract_parser.py — Document processing for Contract AI
Supports PDF (via PyMuPDF) and DOCX (via python-docx)
"""

import io
import re
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
	"""Extract text from PDF bytes using PyMuPDF."""
	try:
		import fitz  # PyMuPDF

		doc = fitz.open(stream=file_bytes, filetype="pdf")
		text_parts = []
		for page_num, page in enumerate(doc):
			text = page.get_text("text")
			if text.strip():
				text_parts.append(f"[Page {page_num + 1}]\n{text}")
		doc.close()
		return "\n\n".join(text_parts)
	except Exception as e:
		raise ValueError(f"Failed to extract PDF text: {str(e)}") from e


def extract_text_from_docx(file_bytes: bytes) -> str:
	"""Extract text from DOCX bytes using python-docx."""
	try:
		from docx import Document

		doc = Document(io.BytesIO(file_bytes))
		text_parts = []
		for para in doc.paragraphs:
			if para.text.strip():
				text_parts.append(para.text)
		for table in doc.tables:
			for row in table.rows:
				row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
				if row_text:
					text_parts.append(row_text)
		return "\n".join(text_parts)
	except Exception as e:
		raise ValueError(f"Failed to extract DOCX text: {str(e)}") from e


def extract_text(file_bytes: bytes, filename: str) -> str:
	"""Auto-detect file type and extract text."""
	filename_lower = filename.lower()
	if filename_lower.endswith(".pdf"):
		return extract_text_from_pdf(file_bytes)
	if filename_lower.endswith(".docx"):
		return extract_text_from_docx(file_bytes)
	if filename_lower.endswith(".doc"):
		raise ValueError("Legacy .doc format not supported. Please convert to .docx")
	try:
		return file_bytes.decode("utf-8")
	except Exception as exc:
		raise ValueError(f"Unsupported file format: {filename}") from exc


def clean_text(text: str) -> str:
	"""Clean and normalize extracted text."""
	text = re.sub(r"\n{3,}", "\n\n", text)
	text = re.sub(r" {2,}", " ", text)
	text = text.replace("\x00", "")
	lines = [line.strip() for line in text.split("\n")]
	return "\n".join(lines).strip()


def get_contract_summary_for_context(text: str, max_length: int = 12000) -> str:
	"""Truncated contract text suitable for model context."""
	if len(text) <= max_length:
		return text
	first_portion = int(max_length * 0.66)
	last_portion = max_length - first_portion
	first_part = text[:first_portion]
	last_part = text[-last_portion:]
	return (
		first_part
		+ f"\n\n[... CONTRACT CONTINUES — {len(text) - max_length} characters omitted for brevity ...]\n\n"
		+ last_part
	)
